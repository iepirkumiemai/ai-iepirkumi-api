import base64
from io import BytesIO
from typing import List, Optional

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import JSONResponse
from docx import Document

from dropbox_client import get_dropbox

app = FastAPI(title="AI Iepirkumi API")


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/ai-tender/analyze")
async def analyze_tender(
    candidate_name: str = Form(...),

    # 1. variants – augšupielādēti faili
    tender_file: Optional[UploadFile] = File(None),
    candidate_archive: Optional[UploadFile] = File(None),

    # 2. variants – ceļi Dropbox mapē
    tender_dropbox_path: Optional[str] = Form(None),
    candidate_dropbox_path: Optional[str] = Form(None),
):
    """
    Minimāls endpoints:
    - Ja ir tender_file/candidate_archive -> izmanto tos
    - Citādi, ja ir tender_dropbox_path/candidate_dropbox_path -> ņem no Dropbox
    - Atgriež testu analīzi + DOCX base64
    """

    tender_bytes: Optional[bytes] = None
    candidate_bytes: Optional[bytes] = None

    # Nolikums
    if tender_file is not None:
        tender_bytes = await tender_file.read()
    elif tender_dropbox_path:
        dbx = get_dropbox()
        if not dbx.file_exists(tender_dropbox_path):
            raise HTTPException(status_code=400, detail="Tender file not found in Dropbox")
        tender_bytes = dbx.download_file(tender_dropbox_path)

    # Kandidāta arhīvs
    if candidate_archive is not None:
        candidate_bytes = await candidate_archive.read()
    elif candidate_dropbox_path:
        dbx = get_dropbox()
        if not dbx.file_exists(candidate_dropbox_path):
            raise HTTPException(status_code=400, detail="Candidate archive not found in Dropbox")
        candidate_bytes = dbx.download_file(candidate_dropbox_path)

    if tender_bytes is None or candidate_bytes is None:
        raise HTTPException(
            status_code=400,
            detail="Provide either uploaded files (tender_file, candidate_archive) or Dropbox paths (tender_dropbox_path, candidate_dropbox_path).",
        )

    # ŠEIT vēlāk būs īstā OpenAI analīze
    short_conclusion = (
        f"TESTA ANALĪZE: kandidāts '{candidate_name}' ir apstrādāts "
        f"(šobrīd bez īstas AI analīzes – tikai skeletons)."
    )

    detailed_analysis = (
        "Šis ir pagaidu detalizētās analīzes teksts. "
        "Vēlāk šeit ievietosim reālo OpenAI ģenerēto analīzi pēc nolikuma prasībām."
    )

    table_html = """
    <table border="1" cellpadding="4" cellspacing="0">
      <thead>
        <tr>
          <th>Nr.</th>
          <th>Punkts</th>
          <th>Prasība</th>
          <th>Pamatojums</th>
          <th>Rezultāts</th>
        </tr>
      </thead>
      <tbody>
        <tr>
          <td>1</td>
          <td>TEST-1</td>
          <td>Testa prasība</td>
          <td>Šī ir tikai testa rinda, lai pārbaudītu integrāciju.</td>
          <td>ATBILST (TESTS)</td>
        </tr>
      </tbody>
    </table>
    """

    error_flags: List[dict] = []

    doc = Document()
    doc.add_heading("Pretendenta kvalifikācijas izvērtējums (TESTS)", level=1)
    doc.add_paragraph(f"Kandidāts: {candidate_name}")

    doc.add_heading("1. Īsais slēdziens", level=2)
    doc.add_paragraph(short_conclusion)

    doc.add_heading("2. Detalizēta analīze (TESTA SKELETONS)", level=2)
    doc.add_paragraph(detailed_analysis)

    doc.add_heading("3. Kopsavilkuma tabula (TESTA)", level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Nr."
    hdr[1].text = "Punkts"
    hdr[2].text = "Prasība"
    hdr[3].text = "Pamatojums"
    hdr[4].text = "Rezultāts"

    row = table.add_row().cells
    row[0].text = "1"
    row[1].text = "TEST-1"
    row[2].text = "Testa prasība"
    row[3].text = "Šī ir tikai testa rinda, lai pārbaudītu integrāciju."
    row[4].text = "ATBILST (TESTS)"

    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")

    response = {
        "short_conclusion": short_conclusion,
        "detailed_analysis": detailed_analysis,
        "table_html": table_html,
        "error_flags": error_flags,
        "docx_base64": docx_b64,
        "pdf_base64": None,
    }

    return JSONResponse(response)
